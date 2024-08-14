from fpdf import FPDF
import pandas as pd
from docx import Document
from io import BytesIO

def dataframe_to_pdf(df, output_buffer):
    """
    Bir DataFrame'i PDF formatına dönüştürür ve belirtilen dosya yoluna kaydeder.

    Args:
        df (pd.DataFrame): PDF'ye dönüştürülecek DataFrame.
        output_buffer (BytesIO): PDF'nin yazılacağı buffer.

    Returns:
        None
    """
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    # Başlık ekleme (isteğe bağlı)
    pdf.cell(200, 10, txt="Sentiment Analysis Results", ln=True, align='C')
    
    # DataFrame sütun başlıklarını ekleme
    for col in df.columns:
        pdf.cell(40, 10, col, border=1)
    pdf.ln()

    # DataFrame satırlarını ekleme
    for index, row in df.iterrows():
        for col in df.columns:
            pdf.cell(40, 10, str(row[col]), border=1)
        pdf.ln()
    
    # PDF'yi buffer'a kaydet
    pdf.output(output_buffer)

def dataframe_to_docx(df, output_buffer):
    """
    Bir DataFrame'i DOCX formatına dönüştürür ve belirtilen dosya yoluna kaydeder.

    Args:
        df (pd.DataFrame): DOCX'e dönüştürülecek DataFrame.
        output_buffer (BytesIO): DOCX'in yazılacağı buffer.

    Returns:
        None
    """
    doc = Document()
    
    # Başlık ekleme (isteğe bağlı)
    doc.add_heading('Sentiment Analysis Results', level=1)
    
    # DataFrame'i DOCX tablosuna dönüştürme
    table = doc.add_table(rows=1, cols=len(df.columns))
    
    # Sütun başlıklarını ekleme
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    
    # Satırları ekleme
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            row_cells[i].text = str(row[col])
    
    # DOCX dosyasını buffer'a kaydetme
    doc.save(output_buffer)
